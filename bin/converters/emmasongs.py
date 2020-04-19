from docx import Document
from docx.shared import Cm, Pt
import collections
import os

from converters.base import AbstractConverter
from converters.features.chords import GuitarChordsWithoutPositioning


class EmmaSongsConverter(AbstractConverter):
    def __init__(self, args):
        super().__init__()
        self._preprocessor.set_required_features(GuitarChordsWithoutPositioning())
        self._from_dir = args.from_dir
        self._to_dir = args.to_dir

    @staticmethod
    def create_argparser(subparsers):
        parser_emmasongs = subparsers.add_parser("emmasongs", help="Converts to the EmmaSongs app's format.")
        parser_emmasongs.add_argument("--to-dir", required=True, help="directory of target files; will be deleted if exists")
        return parser_emmasongs

    def setup(self):
        self._create_out_dir(os.path.join(self._to_dir, 'docx'))
        self._create_out_dir(os.path.join(self._to_dir, 'txt'))

    def convert(self, song_yaml, filepath):
        chords = song_yaml.get('chords')
        self._preprocessor.preprocess(song_yaml, soft_line_break_strategy='ignore')

        # Find Hungarian book, check if E* or I*
        hun_book = self._get_book_from_yaml(song_yaml, 'emm_hu')
        if hun_book is None \
                or hun_book['number'].startswith('E') or hun_book['number'].startswith('I'):
            return

        # Skip if song number is 7## (alleluia) and French book doesn't exist
        if hun_book['number'].startswith('7') and len(hun_book['number']) == 3:
            fr_book = self._get_book_from_yaml(song_yaml, 'emm_fr')
            if fr_book is None:
                return

        # Find Hungarian lyrics
        hun_lyrics = self._get_lyrics_from_yaml(song_yaml, 'hu')
        if hun_lyrics is None:
            return

        document = Document()
        document.styles['Normal'].font.name = 'Arial'
        document.styles['Normal'].font.size = Pt(10)

        self._add_header(document, song_yaml, hun_lyrics)
        self._add_verses(document, hun_lyrics, chords)
        self._add_copyright_info(document, song_yaml, hun_lyrics)

        out_filename_stub = os.path.splitext(os.path.basename(filepath))[0]

        document.save(os.path.join(self._to_dir, 'docx', out_filename_stub+".docx"))

        doc_text = [p.text for p in document.paragraphs]
        with open(os.path.join(self._to_dir, 'txt', out_filename_stub+".txt"), 'wt') as f:
            f.write('\n\n'.join(doc_text))

    @staticmethod
    def _set_tabstops(p):
        p.paragraph_format.tab_stops.add_tab_stop(Cm(10.0))

    def _add_header(self, document, song_yaml, hun_lyrics):
        p = document.add_paragraph()
        self._set_tabstops(p)

        # Add title line
        fr_book = self._get_book_from_yaml(song_yaml, 'emm_fr')
        if fr_book is None:
            fr_songno = ""
        else:
            fr_songno = fr_book['number'] + "-HU"

        r = p.add_run(hun_lyrics['title'] + "\t" + fr_songno)
        r.font.size = Pt(12)
        r.font.bold = True

        # Add about section
        if 'about' not in song_yaml:
            return
        music_by = song_yaml['about'].get('music')
        lyrics_by = song_yaml['about'].get('lyrics')
        if music_by is None and lyrics_by is None:
            return
        if music_by is not None and lyrics_by is not None and music_by == lyrics_by:
            r = p.add_run("\nMusic and lyrics: "+music_by)
        else:
            ml_text = []
            if music_by is not None:
                ml_text.append("Music: "+music_by)
            if lyrics_by is not None:
                ml_text.append("Lyrics: "+lyrics_by)
            r = p.add_run("\n" + "\n".join(ml_text))
        r.font.size = Pt(8)

    def _add_verses(self, document, hun_lyrics, chords):
        # Preprocess chords
        if chords is not None:
            chord_map = {c['template']: c for c in chords}
            for template in chords:
                if 'verses' not in template:
                    continue
                for template_verse in template['verses']:
                    chord_map[template_verse] = template
                del chord_map[template['template']]

        for verse in hun_lyrics['verses']:
            p = document.add_paragraph()
            self._set_tabstops(p)

            # Find out displayed verse name
            verse_name = self._get_printed_verse_name(verse['name'].upper())

            # Find chord template
            verse_chords_lines = None
            used_chord_template = None
            if chords is not None:
                if verse['name'] in chord_map:
                    used_chord_template = verse['name']
                elif verse['name'][0] in chord_map:
                    used_chord_template = verse['name'][0]
                elif 'all' in chord_map:
                    used_chord_template = 'all'
                if used_chord_template is not None:
                    verse_chords = chord_map[used_chord_template]
                    del chord_map[used_chord_template]  # So that we print it only once
                    verse_chords_lines = [l for l in verse_chords['lines'] if l is not None]

            # Compile verse header
            r = p.add_run(verse_name)
            self._bold_if_chorus(verse, r)
            if verse_chords_lines is not None:
                r = p.add_run("\tChords for " + self._get_printed_chord_template_name(used_chord_template))
                r.font.size = Pt(8)
                r.font.italic = True

            # Print verse lines
            verse_lines = [line for line in verse['lines'] if line is not None and line != '']
            printed_lines = self._compile_verse(verse_lines, verse_chords_lines)

            r = p.add_run("\n" + "\n".join(printed_lines))
            self._bold_if_chorus(verse, r)

    def _compile_verse(self, lines, chords):
        printed_lines = []
        for i, line in enumerate(lines):
            line_chords = chords[i] if chords is not None else None
            if isinstance(line, collections.Mapping):
                # This is a repeat group.
                printed_group_lines = []
                for gi, group_line in enumerate(line['lines']):
                    # First, format the line: add the repeat signs.
                    printed_line = ''
                    if gi == 0:
                        printed_line += '/: '
                    printed_line += group_line
                    if gi == len(line['lines'])-1:
                        printed_line += ' :/'

                    # Second, add the chords.
                    if line_chords is not None:
                        # If all repetitions are uniform, then we can just copy the chords as normal.
                        if line_chords['type'] == 'uniform':
                            printed_line += '\t' + ' '.join(self._remove_placeholders_from_chord_list(line_chords['lines'][gi]))
                        # Otherwise, we need to specify all iterations.
                        elif line_chords['type'] == 'unique':
                            # First, let's check if all repetitions have the same chords for this line,
                            # because then we don't need to list them all.
                            if self._list_elems_all_same(c[gi] for c in line_chords['repetitions']):
                                printed_line += '\t' + ' '.join(self._remove_placeholders_from_chord_list(line_chords['repetitions'][0][gi]))
                            else:
                                # They're different, so we'll need to list them on separate lines, with (1), (2), ... prefixes for the repetitions.
                                repetition_chords = []
                                for rid, repetition in enumerate(line_chords['repetitions']):
                                    repetition_chords.append("\t(" + str(rid+1) + ") " + ' '.join(self._remove_placeholders_from_chord_list(repetition[gi])))
                                printed_line += '\n'.join(repetition_chords)
                        else:
                            raise Exception("Unknown repeat type: "+line_chords['type'])
                    printed_group_lines.append(printed_line)
                printed_lines.append("\n".join(printed_group_lines))
            else:
                # This is a simple line
                printed_line = line
                if line_chords is not None:
                    printed_line += "\t" + ' '.join(self._remove_placeholders_from_chord_list(line_chords))
                printed_lines.append(printed_line)
        return printed_lines

    @staticmethod
    def _remove_placeholders_from_chord_list(chord_list):
        return (c for c in chord_list if c != '_')

    @staticmethod
    def _get_printed_verse_name(verse_name):
        if verse_name.startswith('C'):
            return 'Ch.'
        elif verse_name.startswith('V'):
            return verse_name[1:] + '.'
        elif verse_name.startswith('B'):
            return 'Bridge'
        else:
            return verse_name

    @staticmethod
    def _get_printed_chord_template_name(chord_template_name):
        if chord_template_name == 'all':
            return "all verses"
        elif chord_template_name == 'c':
            return "all choruses"
        elif chord_template_name == 'v':
            return "all numbered verses"
        elif chord_template_name == 'b':
            return "all bridges"
        elif chord_template_name.startswith('c'):
            return "chorus " + chord_template_name[1:]
        elif chord_template_name.startswith('v'):
            return "verse " + chord_template_name[1:]
        elif chord_template_name.startswith('b'):
            return "bridge " + chord_template_name[1:]
        else:
            raise Exception("Unknown chord template name: " + chord_template_name)

    @staticmethod
    def _bold_if_chorus(verse, run):
        if verse['name'].startswith("c"):
            run.font.bold = True

    @staticmethod
    def _list_elems_all_same(li):
        element = None
        for item in li:
            if element is None:
                element = item
            elif element != item:
                return False
        return True

    def _add_copyright_info(self, document, song_yaml, hun_lyrics):
        copr_lines = []

        # Original copyright
        if 'about' in song_yaml:
            orig_copr_line = ''
            if 'c_year' in song_yaml['about']:
                orig_copr_line += str(song_yaml['about']['c_year']) + ', '
            if 'c_holder' in song_yaml['about']:
                orig_copr_line += song_yaml['about']['c_holder']
            if orig_copr_line != '':
                copr_lines.append('© '+orig_copr_line)

        # Hungarian copyright
        if 'about' in hun_lyrics:
            hun_copr_line = ''
            if 'adapted_by' in hun_lyrics['about']:
                hun_copr_line += 'Adapted by: ' + hun_lyrics['about']['adapted_by'] + ' '
            if 'c_year' in hun_lyrics['about'] or 'c_holder' in hun_lyrics['about']:
                hun_copr_line += '© '
            if 'c_year' in hun_lyrics['about']:
                hun_copr_line += str(hun_lyrics['about']['c_year']) + ', '
            if 'c_holder' in hun_lyrics['about']:
                hun_copr_line += hun_lyrics['about']['c_holder']
            if hun_lyrics != '':
                copr_lines.append(hun_copr_line)

        if copr_lines:
            p = document.add_paragraph()
            self._set_tabstops(p)
            r = p.add_run('\n'.join(copr_lines))
            r.font.size = Pt(8)
